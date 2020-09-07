#! /opt/local/bin/python
# coding=utf-8
from docx import Document
import re

doc = Document('doc_20200904.docx')
i = 0
bStart = False
bEnd = False
nStart = -1
nEnd = -1
subDoc = Document()
name =""
for x in doc.paragraphs:
    if( not bStart and x.style.name == 'Heading 4' and x.text[:13] == 'Rule_E_Score_'):
        bStart = True
        nStart = i
        print("start")
        print( i)
        print(x.style.name)
        print(x.text)
        name = x.text
    elif( bStart and (x.style.name == 'Heading 4' and x.text[:13] == 'Rule_E_Score_') ):
        bStart = True
        bEnd = True
        nEnd = i
        print("end & start")
        print( i)
        print(x.style.name)
        print(x.text)
        name = x.text
    elif( bStart and (( x.style.name == 'Heading 1') or ( x.style.name == 'Heading 2') or ( x.style.name == 'Heading 3'))):
        bStart = False
        bEnd = True
        nEnd = i
        print("end")
        print( i)
        print(x.style.name)
        print(x.text)
    if( bStart ):
        subDoc.add_paragraph(x)
    if( bEnd ):
        subDoc.save("files/"+name)
        subDoc = Document()
        name =""
        bEnd = False
    i += 1
